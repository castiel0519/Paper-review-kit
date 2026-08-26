#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_docx.py — 依据 summaries/summary_*.json 生成中英双语读书报告 DOCX。
输出：deliverables/机器学习与人工智能在微流控中的应用_读书报告.docx
"""
import json
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
BASE = os.path.dirname(HERE)
SUM_DIR = os.path.join(BASE, "summaries")
FIG_DIR = os.path.join(BASE, "papers_figs")
DELIVER_DIR = os.path.join(BASE, "deliverables")
OUT = os.path.join(DELIVER_DIR, "机器学习与人工智能在微流控中的应用_读书报告.docx")

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)


def set_cn(run, name="微软雅黑"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=None,
             align=None, space_after=6, first_line_indent=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        set_cn(r)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.size = Pt({1: 22, 2: 16, 3: 13, 4: 11.5}.get(level, 12))
        r.font.color.rgb = NAVY if level <= 2 else ACCENT
        r.font.bold = True
        set_cn(r)
    return h


def add_meta_table(doc, s):
    rows = [
        ("英文标题", s.get("title_en", "")),
        ("中文标题", s.get("title_zh", "")),
        ("期刊 / 年份", f"{s.get('journal', '')} · {s.get('year', '')}"),
        ("DOI / PMID / PMCID", " / ".join([x for x in [s.get("doi"), s.get("pmid"), s.get("pmcid")] if x]) or "N/A"),
        ("研究方向", s.get("kind_zh", "")),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        for j, txt in enumerate([k, str(v)]):
            cell = t.cell(i, j)
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(9.5)
            r.font.bold = (j == 0)
            set_cn(r)
        t.cell(i, 0).width = Inches(1.5)
        t.cell(i, 1).width = Inches(5.6)
    doc.add_paragraph()
    return t


def add_paper(doc, s, scheme_path):
    add_heading(doc, f"论文 {s['id']} · {s.get('title_zh', '')}", level=2)
    add_meta_table(doc, s)
    if s.get("abstract_en"):
        add_heading(doc, "摘要 (Abstract)", level=3)
        add_para(doc, f"【英文】{s['abstract_en']}", size=10)
        add_para(doc, f"【中文】{s['abstract_zh']}", size=10)
    if s.get("background_zh") or s.get("problem_zh"):
        add_heading(doc, "1. 研究背景与科学问题", level=3)
        if s.get("background_zh"):
            add_para(doc, f"背景：{s['background_zh']}", size=10)
        if s.get("problem_zh"):
            add_para(doc, f"问题：{s['problem_zh']}", size=10)
    if s.get("data_zh") or s.get("task_zh"):
        add_heading(doc, "2. 数据与任务", level=3)
        if s.get("data_zh"):
            add_para(doc, f"数据：{s['data_zh']}", size=10)
        if s.get("task_zh"):
            add_para(doc, f"任务：{s['task_zh']}", size=10)
    add_heading(doc, "3. 方法", level=3)
    add_para(doc, s.get("method_zh", ""), size=10)
    add_heading(doc, "4. 关键结果", level=3)
    add_para(doc, s.get("results_zh", ""), size=10)
    if s.get("metrics_zh"):
        t = doc.add_table(rows=len(s["metrics_zh"]) + 1, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).paragraphs[0].add_run("指标").bold = True
        t.cell(0, 1).paragraphs[0].add_run("数值").bold = True
        for i, m in enumerate(s["metrics_zh"], start=1):
            if isinstance(m, dict):
                k, v = m.get("label", ""), m.get("value", "")
            else:
                k, v = str(m), ""
            c0, c1 = t.cell(i, 0).paragraphs[0], t.cell(i, 1).paragraphs[0]
            r0, r1 = c0.add_run(k), c1.add_run(str(v))
            r0.font.size = r1.font.size = Pt(9.5)
            set_cn(r0); set_cn(r1)
        doc.add_paragraph()
    add_heading(doc, "5. 创新点与局限性", level=3)
    add_para(doc, f"创新点：{s.get('innovation_zh', '')}", size=10)
    add_para(doc, f"局限性：{s.get('limitation_zh', '')}", size=10)
    add_heading(doc, "6. 研究范式 (Research Paradigm)", level=3)
    tags = s.get("paradigm_tags") or []
    add_para(doc, (f"范式标签：{' / '.join(tags)}  " if tags else "") + s.get("paradigm_zh", ""), size=10)
    add_heading(doc, "7. 研究框架解析 (Research Framework)", level=3)
    add_para(doc, s.get("framework_zh", ""), size=10)
    steps = s.get("framework_steps") or []
    if steps:
        for i, st in enumerate(steps):
            add_para(doc, f"步骤 {i+1}：{st}", size=10, space_after=2)
    add_heading(doc, "8. 论文 Scheme 解读", level=3)
    if scheme_path and os.path.exists(scheme_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(scheme_path, width=Inches(5.8))
    add_para(doc, s.get("scheme_zh", ""), size=10)
    add_heading(doc, "9. 对本领域研究的启示", level=3)
    add_para(doc, s.get("lessons_zh", ""), size=10)
    doc.add_paragraph()


def main():
    os.makedirs(DELIVER_DIR, exist_ok=True)
    meta = json.load(open(os.path.join(BASE, "papers_meta.json"), encoding="utf-8"))
    papers = {p["id"]: p for p in meta["papers"]}
    summaries = []
    for fn in sorted(os.listdir(SUM_DIR)):
        if not fn.endswith(".json"):
            continue
        s = json.load(open(os.path.join(SUM_DIR, fn), encoding="utf-8"))
        if s.get("method_zh"):
            summaries.append(s)
    summaries.sort(key=lambda s: s["id"])

    doc = Document()
    # 全局默认字体
    st = doc.styles["Normal"]
    st.font.name = "微软雅黑"
    st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 封面
    add_para(doc, "", size=20)
    add_para(doc, "机器学习与人工智能", size=28, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, "在微流控中的应用", size=28, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_para(doc, "SCI 论文精读调研报告（2020–2025）", size=16, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Machine Learning & Artificial Intelligence in Microfluidics:", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "A Reading Report on Recent SCI Publications (2020–2026)", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, f"覆盖论文数：{len(summaries)} 篇 ｜ 研究方向：数字微流控 / 图像识别 / 智能分选 / 设计优化 / 器官芯片 / 诊断POCT / 单细胞 / 综述",
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "生成日期：2026-08（系统任务自动生成）", size=9,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # 目录式说明
    add_heading(doc, "报告导读", level=1)
    add_para(doc, "本报告基于12篇以上2020–2026年发表的SCI论文精读而成，覆盖机器学习与人工智能在微流控中的八类代表性应用：",
             size=10.5)
    for line in [
        "① 数字微流控（DMF：液滴操控、细胞分选、多态控制）",
        "② 液滴/细胞图像识别与分类（CNN、YOLO、成像流式、无标记液滴分类）",
        "③ 液滴/细胞分选与实时控制（FPGA 加速、无标记并行分选）",
        "④ 实验设计与参数优化（逆向设计、LLM 自主设计、数字孪生）",
        "⑤ 器官芯片/生物传感/药物筛选（digital twin 肝芯片、AI 药物筛选）",
        "⑥ 疾病诊断与 POCT（CTC 全息成像、外泌体 SERS、纸基微流控）",
        "⑦ 单细胞分析与高通量成像（机器人+AI、运动敏感干涉成像）",
        "⑧ 综述与范式（智能微流控、可穿戴集成、AI×微系统）",
    ]:
        add_para(doc, line, size=10.5, space_after=2)
    add_para(doc, "每篇论文按「摘要—背景与问题—数据与任务—方法—关键结果—创新与局限—研究范式—研究框架解析—Scheme解读—启示」十个维度精读整理，"
                  "并附横向对比、范式总结与趋势展望。", size=10.5)

    add_heading(doc, "1. 调研方法", level=2)
    for line in [
        "选文标准：SCI 收录期刊；主题为机器学习/深度学习/人工智能在微流控或细胞图像分析中的应用；优先开放获取(OA)以获得全文PDF；覆盖多维度应用与方法学多样性。",
        "文献检索：结合 PubMed / Europe PMC / Unpaywall / Semantic Scholar / 出版社官网进行检索、元数据核验与PDF获取。",
        "精读流程：下载PDF → 逐页提取全文文本 → 定位摘要、数据、方法、结果、讨论关键段落 → 提取论文Scheme图 → 按统一结构生成中英双语精读摘要。",
        "PDF获取说明：仅对开放获取(OA)论文下载PDF全文；非OA论文保留DOI与元数据并在附录中标注。",
    ]:
        add_para(doc, line, size=10.5, space_after=4)

    add_heading(doc, "2. 论文全景", level=2)
    t = doc.add_table(rows=len(summaries) + 1, cols=5)
    t.style = "Table Grid"
    heads = ["编号", "方向", "论文（中文标题）", "期刊/年份", "方法核心"]
    for j, htext in enumerate(heads):
        r = t.cell(0, j).paragraphs[0].add_run(htext)
        r.font.bold = True; r.font.size = Pt(9.5); set_cn(r)
    for i, s in enumerate(summaries, start=1):
        vals = [s["id"], s.get("kind_zh", ""), s.get("title_zh", ""),
                f"{s.get('journal', '')} {s.get('year', '')}", s.get("method_zh", "")[:28] + ("…" if len(s.get("method_zh", "")) > 28 else "")]
        for j, v in enumerate(vals):
            r = t.cell(i, j).paragraphs[0].add_run(str(v))
            r.font.size = Pt(8.5); set_cn(r)
    doc.add_paragraph()

    add_heading(doc, "3. 论文精读", level=1)
    for s in summaries:
        scheme = os.path.join(FIG_DIR, f"{s['id']}_fig1.png")
        if not os.path.exists(scheme):
            scheme = None
        add_paper(doc, s, scheme)

    add_heading(doc, "4. 横向对比", level=1)
    add_para(doc, "下表汇总各论文的数据规模、任务、主要方法与报告的核心性能指标（仅列出原文报告的代表性指标）。", size=10)
    t = doc.add_table(rows=len(summaries) + 1, cols=6)
    t.style = "Table Grid"
    heads = ["论文", "方向", "数据规模/来源", "任务", "方法", "代表性指标"]
    for j, htext in enumerate(heads):
        r = t.cell(0, j).paragraphs[0].add_run(htext)
        r.font.bold = True; r.font.size = Pt(9); set_cn(r)
    for i, s in enumerate(summaries, start=1):
        m = s.get("metrics_zh") or []
        mstr = ""
        if m:
            m0 = m[0]
            mstr = f"{m0.get('label','')}: {m0.get('value','')}"
        vals = [f"{s['id']} {s.get('kind_zh','')[:6]}", s.get("data_zh", "")[:60],
                s.get("task_zh", "")[:50], s.get("method_zh", "")[:70],
                mstr[:40]]
        # 6列 -> 论文/方向 合并为1列
        vals = [f"{s['id']} {s.get('title_zh', '')[:20]}", s.get("kind_zh", ""),
                s.get("data_zh", "")[:55], s.get("task_zh", "")[:45], s.get("method_zh", "")[:65], mstr[:42]]
        for j, v in enumerate(vals):
            r = t.cell(i, j).paragraphs[0].add_run(str(v))
            r.font.size = Pt(7.5); set_cn(r)
    doc.add_paragraph()

    add_heading(doc, "5. 研究范式总结与趋势", level=1)
    trends = [
        ("范式一：传统ML + 特征工程", "以物理/几何特征（液滴直径、频率、压力、图像纹理、光谱峰）为输入，配合SVM、随机森林、XGBoost、贝叶斯回归；可解释性强、算力低，适用于液滴尺寸预测、生物标志物回归等。"),
        ("范式二：端到端深度学习（CNN/YOLO/TrackNet）", "以液滴/细胞图像、流动显微视频为直接输入，完成检测、分类、分割与追踪；在液滴分类、成像流式、细胞分选中成为主流，精度高但依赖标注与算力。"),
        ("范式三：物理信息/控制驱动学习（RL、digital twin）", "将流体力学物理约束或仿真器/数字孪生嵌入学习环，用深度强化学习、贝叶斯优化、代理模型实现液滴生成、芯片参数、蒸发控制的闭环自适应优化。"),
        ("范式四：生成式与逆向设计（GAN/VAE/扩散模型、LLM）", "由目标功能/图案反向生成芯片结构与实验条件，大语言模型进一步把领域知识文本化，实现自主设计框架与自然语言交互设计。"),
        ("范式五：多模态集成（图像+光谱+组学+临床）", "微流控SERS、无标记定量成像、芯片传感与临床信息融合，实现外泌体肺癌分型、CTC检测、牙周炎纸基诊断等POCT多模态分析。"),
        ("范式六：综述集成与范式化系统", "智能微流控（Matter 2020）、AI×微系统（Micromachines 2023）、机器人+AI单细胞（Lab Chip 2025）等综述构建「数据-算法-硬件-应用」闭环方法学体系。"),
    ]
    for k, v in trends:
        add_para(doc, f"{k}：{v}", size=10, space_after=6)
    add_heading(doc, "趋势与展望", level=2)
    for line in [
        "① 从单一数据集走向多中心、跨设备、跨模态（图像+基因+药理）联合建模。",
        "② 标注瓶颈推动自监督/预训练/基础模型成为标准组件，标注成本有效下降。",
        "③ 从「分类/分割」走向「溯源、功能推断、性质预测」等更高级的推断任务。",
        "④ 临床落地强调可解释性、FDA级验证与大规模人群前瞻性评估（如宫颈细胞学多中心筛查）。",
        "⑤ 微流控正与单细胞组学、空间组学融合，形成「形态-分子」多模态分析范式。",
    ]:
        add_para(doc, line, size=10, space_after=3)

    add_heading(doc, "6. 局限与阅读建议", level=1)
    for line in [
        "① 本报告所选论文以开放获取为主，优先保证PDF全文可获取与精读完整性；非OA论文未纳入精读正文。",
        "② 论文间任务与指标不可直接横向比较（数据集、评价协议、临床场景均不同），阅读时注意区分「方法学严谨性」与「临床可用性」。",
        "③ 建议结合论文补充材料（数据集规模、消融实验、统计分析）进一步验证结论。",
    ]:
        add_para(doc, line, size=10, space_after=3)

    add_heading(doc, "7. 参考文献", level=1)
    for i, s in enumerate(summaries, start=1):
        add_para(doc, f"[{i}] {s.get('title_en','')}. {s.get('journal','')}, {s.get('year','')}. DOI: {s.get('doi','')}",
                 size=9, space_after=2)

    doc.save(OUT)
    print("DOCX written:", OUT, "papers:", len(summaries))


if __name__ == "__main__":
    main()
