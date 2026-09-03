#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_review_docx.py — 将综述 Markdown 转为 DOCX。

用法：
    python scripts/make_review_docx.py                 # 独立输出
    python scripts/make_review_docx.py --append        # 追加到读书报告 DOCX 末尾

输出：deliverables/{title}_调研综述.docx
"""
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from prk_config import load_papers_meta, output_dir, parse_project_arg, project_title


def set_cn(run, name="微软雅黑"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_markdown(doc, text):
    for line in text.splitlines():
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(line.strip())
        for r in (h.runs if 'h' in dir() else doc.paragraphs[-1].runs):
            r.font.size = Pt(10.5)
            set_cn(r)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--review", default="", help="综述 MD 路径")
    parser.add_argument("--append", action="store_true", help="追加到读书报告 DOCX")
    cfg, args = parse_project_arg(parser)
    meta = load_papers_meta(cfg)
    title = project_title(cfg, meta)
    del_dir = output_dir(cfg, "deliverables")
    review_md = Path(args.review) if args.review else del_dir / f"{title}_调研综述.md"
    if not review_md.exists():
        raise SystemExit(f"综述 Markdown 不存在: {review_md}")
    text = review_md.read_text(encoding="utf-8")

    target = del_dir / f"{title}_调研综述.docx"
    if args.append:
        report_doc = del_dir / f"{title}_读书报告.docx"
        if not report_doc.exists():
            raise SystemExit(f"读书报告不存在，无法追加: {report_doc}")
        doc = Document(str(report_doc))
        doc.add_page_break()
        doc.add_heading("调研方向综述", level=1)
    else:
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name = "微软雅黑"
        st.font.size = Pt(10.5)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    add_markdown(doc, text)
    doc.save(str(target))
    print("review docx ->", target, "append=", args.append)


if __name__ == "__main__":
    main()
