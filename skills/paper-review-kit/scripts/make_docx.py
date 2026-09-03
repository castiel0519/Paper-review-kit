#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_docx.py — 依据 summaries/summary_*.json 生成中英双语读书报告 DOCX。
标题、年份范围、文件名、叙事文案由 project.yml 驱动；未配置时使用通用兜底。
输出：deliverables/{title}_读书报告.docx
"""
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from prk_config import (
    cfg_get, load_papers_meta, output_dir, parse_project_arg, project_path, project_range,
    project_title, read_json,
)
from prk_schema import validate_summary

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)


def set_cn(run, name="微软雅黑"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=None,
             align=None, space_after=6, first_line_indent=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        set_cn(r)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.size = Pt({1: 22, 2: 16, 3: 13, 4: 11.5}.get(level, 12))
        r.font.color.rgb = NAVY if level <= 2 else ACCENT
        r.font.bold = True
        set_cn(r)
    return h


def add_meta_table(doc, s):
    rows = [
        ("英文标题", s.get("title_en", "")),
        ("中文标题", s.get("title_zh", "")),
        ("期刊 / 年份", f"{s.get('journal', '')} · {s.get('year', '')}"),
        ("DOI / PMID / PMCID", " / ".join([str(x) for x in [s.get("doi"), s.get("pmid"), s.get("pmcid")]
                                            if x is not None and str(x).strip()]) or "N/A"),
        ("研究方向", s.get("kind_zh", "")),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        for j, txt in enumerate([k, str(v)]):
            cell = t.cell(i, j)
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(9.5)
            r.font.bold = (j == 0)
            set_cn(r)
        t.cell(i, 0).width = Inches(1.5)
        t.cell(i, 1).width = Inches(5.6)
    doc.add_paragraph()
    return t


def add_paper(doc, s, scheme_path):
    add_heading(doc, f"论文 {s['id']} · {s.get('title_zh', '')}", level=2)
    if s.get("pdf_status") in ("abstract_only", "missing"):
        label = "仅摘要" if s.get("pdf_status") == "abstract_only" else "全文缺失"
        add_para(doc, f"全文状态：{label}（{s.get('pdf_status')}）。以下内容基于可获取信息整理，不编造全文细节。",
                 size=9.5, color=ACCENT, space_after=6)
    add_meta_table(doc, s)
    if s.get("abstract_en"):
        add_heading(doc, "摘要 (Abstract)", level=3)
        add_para(doc, f"【英文】{s['abstract_en']}", size=10)
        add_para(doc, f"【中文】{s['abstract_zh']}", size=10)
    if s.get("background_zh") or s.get("problem_zh"):
        add_heading(doc, "1. 研究背景与科学问题", level=3)
        if s.get("background_zh"):
            add_para(doc, f"背景：{s['background_zh']}", size=10)
        if s.get("problem_zh"):
            add_para(doc, f"问题：{s['problem_zh']}", size=10)
    if s.get("data_zh") or s.get("task_zh"):
        add_heading(doc, "2. 数据与任务", level=3)
        if s.get("data_zh"):
            add_para(doc, f"数据：{s['data_zh']}", size=10)
        if s.get("task_zh"):
            add_para(doc, f"任务：{s['task_zh']}", size=10)
    add_heading(doc, "3. 方法", level=3)
    add_para(doc, s.get("method_zh", ""), size=10)
    add_heading(doc, "4. 关键结果", level=3)
    add_para(doc, s.get("results_zh", ""), size=10)
    if s.get("metrics_zh"):
        t = doc.add_table(rows=len(s["metrics_zh"]) + 1, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).paragraphs[0].add_run("指标").bold = True
        t.cell(0, 1).paragraphs[0].add_run("数值").bold = True
        for i, m in enumerate(s["metrics_zh"], start=1):
            if isinstance(m, dict):
                k, v = m.get("label", ""), m.get("value", "")
            else:
                k, v = str(m), ""
            c0, c1 = t.cell(i, 0).paragraphs[0], t.cell(i, 1).paragraphs[0]
            r0, r1 = c0.add_run(k), c1.add_run(str(v))
            r0.font.size = r1.font.size = Pt(9.5)
            set_cn(r0); set_cn(r1)
        doc.add_paragraph()
    add_heading(doc, "5. 创新点与局限性", level=3)
    add_para(doc, f"创新点：{s.get('innovation_zh', '')}", size=10)
    add_para(doc, f"局限性：{s.get('limitation_zh', '')}", size=10)
    add_heading(doc, "6. 研究范式 (Research Paradigm)", level=3)
    tags = s.get("paradigm_tags") or []
    add_para(doc, (f"范式标签：{' / '.join(tags)}  " if tags else "") + s.get("paradigm_zh", ""), size=10)
    add_heading(doc, "7. 研究框架解析 (Research Framework)", level=3)
    add_para(doc, s.get("framework_zh", ""), size=10)
    steps = s.get("framework_steps") or []
    if steps:
        for i, st in enumerate(steps):
            add_para(doc, f"步骤 {i+1}：{st}", size=10, space_after=2)
    add_heading(doc, "8. 论文 Scheme 解读", level=3)
    if scheme_path and os.path.exists(scheme_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(scheme_path), width=Inches(5.8))
    add_para(doc, s.get("scheme_zh", ""), size=10)
    add_heading(doc, "9. 对本领域研究的启示", level=3)
    add_para(doc, s.get("lessons_zh", ""), size=10)
    doc.add_paragraph()


def _report_lines(cfg, key, default):
    val = cfg_get(cfg, "report", "narrative", key)
    if isinstance(val, list) and val:
        return [str(x) for x in val]
    if isinstance(val, str) and val.strip():
        return [val]
    return default


def _generic_trends(cfg):
    return _report_lines(cfg, "trends", [
        "数据驱动建模：以可复用数据集与基准测试为底座，形成可比较、可复现的实验体系。",
        "端到端深度学习：以图像、序列或多模态信号为直接输入，完成检测、分类、分割与预测。",
        "物理信息与控制驱动学习：将领域先验、仿真器或闭环控制嵌入学习过程，提升样本效率与可解释性。",
        "生成式与自主实验：由目标反向设计结构与条件，大语言模型/智能体逐步参与实验设计与优化。",
        "多模态集成：融合图像、光谱、组学与临床信息，推动从研究原型到应用落地。",
    ])


def _generic_outlook(cfg):
    return _report_lines(cfg, "outlook", [
        "从单一数据集走向多中心、跨设备、跨模态联合建模。",
        "标注瓶颈推动自监督/预训练/基础模型成为标准组件。",
        "从「分类/分割」走向「溯源、功能推断、性质预测」等更高级任务。",
        "应用落地强调可解释性、可重复性与前瞻性验证。",
        "方法学综述将推动「数据-算法-硬件-应用」闭环范式化。",
    ])


def _resolve_scheme(s, fig_dir, cfg=None):
    """优先使用 summary.scheme_image，缺失时回退到纸张默认 fig1。"""
    raw = s.get("scheme_image") or ""
    if raw:
        p = Path(raw)
        if not p.is_absolute():
            # 兼容两种常见写法：papers_figs/xx.png 或 xx.png
            try:
                if cfg:
                    root_rel = Path(cfg.get("_project_root", ".")) / p
                    if root_rel.exists():
                        return root_rel
            except Exception:
                pass
            p = fig_dir / p
        if p.exists():
            return p
    default = fig_dir / f"{s.get('id', '')}_fig1.png"
    return default if default.exists() else None


def main():
    cfg, args = parse_project_arg()
    meta = load_papers_meta(cfg)
    title = project_title(cfg, meta)
    rng = project_range(cfg, meta)
    topic_en = cfg_get(cfg, "project", "topic_en", default="Paper Review")
    subtitle = cfg_get(cfg, "report", "cover_subtitle", default="SCI 论文精读调研报告")

    sum_dir = output_dir(cfg, "summaries")
    fig_dir = output_dir(cfg, "papers_figs")
    del_dir = output_dir(cfg, "deliverables")

    summaries = []
    if sum_dir.is_dir():
        for fn in sorted(sum_dir.iterdir()):
            if fn.suffix != ".json":
                continue
            s = read_json(fn, default=None)
            if not isinstance(s, dict) or not s.get("id"):
                continue
            errors, _ = validate_summary(s)
            if errors:
                print(f"  [warn] {fn.name}: " + "; ".join(errors))
            summaries.append(s)
    summaries.sort(key=lambda s: str(s["id"]))

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "微软雅黑"
    st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 封面：标题全部来自配置 / papers_meta
    add_para(doc, "", size=20)
    add_para(doc, title, size=26, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_para(doc, f"{subtitle}（{rng}）", size=16, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, f"{topic_en}:", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, f"A Reading Report on Recent SCI Publications ({rng})", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    kinds = []
    for s in summaries:
        k = s.get("kind_zh") or s.get("kind") or ""
        if k and k not in kinds:
            kinds.append(str(k))
    kind_text = " / ".join(kinds) if kinds else "论文调研"
    add_para(doc, f"覆盖论文数：{len(summaries)} 篇 ｜ 研究方向：{kind_text}",
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "生成日期：系统任务自动生成", size=9,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # 目录式说明
    add_heading(doc, "报告导读", level=1)
    add_para(doc, f"本报告基于 {rng} 年发表的 SCI 论文精读整理而成，"
                  f"共覆盖 {len(kinds)} 类研究方向：{kind_text}。", size=10.5)
    add_para(doc, "每篇论文按「摘要—背景与问题—数据与任务—方法—关键结果—创新与局限—"
                  "研究范式—研究框架解析—Scheme解读—启示」十个维度精读整理，"
                  "并附横向对比、范式总结与趋势展望。", size=10.5)

    add_heading(doc, "1. 调研方法", level=1)
    for line in _report_lines(cfg, "methods", [
        "选文标准：SCI 收录期刊；主题聚焦；优先开放获取(OA)以获得全文PDF；覆盖多维度应用与方法学多样性。",
        "文献检索：结合 PubMed / Europe PMC / Crossref / OpenAlex / Unpaywall 与出版社官网进行检索、元数据核验与PDF获取。",
        "精读流程：下载PDF → 逐页提取全文文本 → 定位摘要、数据、方法、结果、讨论关键段落 → 提取论文Scheme图 → 按统一结构生成中英双语精读摘要。",
        "PDF获取说明：仅对开放获取(OA)及用户明确授权的来源下载PDF全文；非OA论文保留DOI与元数据并标注状态。",
    ]):
        add_para(doc, line, size=10.5, space_after=4)

    add_heading(doc, "2. 论文全景", level=1)
    t = doc.add_table(rows=len(summaries) + 1, cols=5)
    t.style = "Table Grid"
    heads = ["编号", "方向", "论文（中文标题）", "期刊/年份", "方法核心"]
    for j, htext in enumerate(heads):
        r = t.cell(0, j).paragraphs[0].add_run(htext)
        r.font.bold = True; r.font.size = Pt(9.5); set_cn(r)
    for i, s in enumerate(summaries, start=1):
        method = str(s.get("method_zh", "") or "")
        vals = [s["id"], s.get("kind_zh", ""), s.get("title_zh", ""),
                f"{s.get('journal', '')} {s.get('year', '')}",
                method[:28] + ("…" if len(method) > 28 else "")]
        for j, v in enumerate(vals):
            r = t.cell(i, j).paragraphs[0].add_run(str(v))
            r.font.size = Pt(8.5); set_cn(r)
    doc.add_paragraph()

    add_heading(doc, "3. 论文精读", level=1)
    for s in summaries:
        scheme = _resolve_scheme(s, fig_dir, cfg)
        add_paper(doc, s, scheme)

    add_heading(doc, "4. 横向对比", level=1)
    add_para(doc, "下表汇总各论文的数据规模、任务、主要方法与报告的核心性能指标（仅列出原文报告的代表性指标）。", size=10)
    t = doc.add_table(rows=len(summaries) + 1, cols=6)
    t.style = "Table Grid"
    heads = ["论文", "方向", "数据规模/来源", "任务", "方法", "代表性指标"]
    for j, htext in enumerate(heads):
        r = t.cell(0, j).paragraphs[0].add_run(htext)
        r.font.bold = True; r.font.size = Pt(9); set_cn(r)
    for i, s in enumerate(summaries, start=1):
        m = s.get("metrics_zh") or []
        mstr = ""
        if m:
            m0 = m[0]
            label = m0.get("label", "") if isinstance(m0, dict) else str(m0)
            value = m0.get("value", "") if isinstance(m0, dict) else ""
            mstr = f"{label}: {value}"
        vals = [f"{s['id']} {str(s.get('title_zh', ''))[:20]}", s.get("kind_zh", ""),
                str(s.get("data_zh", ""))[:55], str(s.get("task_zh", ""))[:45],
                str(s.get("method_zh", ""))[:65], mstr[:42]]
        for j, v in enumerate(vals):
            r = t.cell(i, j).paragraphs[0].add_run(str(v))
            r.font.size = Pt(7.5); set_cn(r)
    doc.add_paragraph()

    add_heading(doc, "5. 研究范式总结与趋势", level=1)
    for line in _generic_trends(cfg):
        add_para(doc, line, size=10, space_after=6)
    add_heading(doc, "趋势与展望", level=2)
    for line in _generic_outlook(cfg):
        add_para(doc, line, size=10, space_after=3)

    add_heading(doc, "6. 局限与阅读建议", level=1)
    for line in _report_lines(cfg, "limitations", [
        "本报告优先保证PDF全文可获取与精读完整性；非OA论文以摘要和元数据参与分析，并显式标注。",
        "论文间任务与指标不可直接横向比较（数据集、评价协议、应用场景均不同），阅读时注意区分「方法学严谨性」与「应用可用性」。",
        "建议结合论文补充材料（数据集规模、消融实验、统计分析）进一步验证结论。",
    ]):
        add_para(doc, line, size=10, space_after=3)

    add_heading(doc, "7. 参考文献", level=1)
    for i, s in enumerate(summaries, start=1):
        add_para(doc, f"[{i}] {s.get('title_en','')}. {s.get('journal','')}, {s.get('year','')}. DOI: {s.get('doi','')}",
                 size=9, space_after=2)

    out = del_dir / f"{title}_读书报告.docx"
    doc.save(str(out))
    print("DOCX written:", out, "papers:", len(summaries))


if __name__ == "__main__":
    main()

